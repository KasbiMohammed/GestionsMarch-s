"""
Export Word
Génération de documents Word pour les rapports et exports
"""

from datetime import datetime
from typing import List, Dict, Optional
from sqlalchemy.orm import Session
from io import BytesIO

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from app.models.market import Market
from app.models.stage import Stage


class WordExportService:
    """Service pour la génération de documents Word"""
    
    def __init__(self, db: Session):
        self.db = db
    
    def create_market_report(self, market_id: int) -> BytesIO:
        """
        Crée un rapport Word détaillé pour un marché
        
        Args:
            market_id: ID du marché
            
        Returns:
            Fichier Word en mémoire
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx n'est pas installé. Installez-le avec: pip install python-docx")
        
        market = self.db.query(Market).filter(Market.id == market_id).first()
        if not market:
            raise ValueError("Marché non trouvé")
        
        doc = Document()
        
        # Titre du document
        title = doc.add_heading('Rapport de Marché Public', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Informations générales
        doc.add_heading('Informations Générales', level=1)
        
        info_table = doc.add_table(rows=1, cols=2)
        info_table.style = 'Table Grid'
        
        info_data = [
            ('Numéro de marché', market.market_number or '-'),
            ('Objet', market.object or '-'),
            ('Maître d\'ouvrage', market.owner or '-'),
            ('Type', market.type.value if market.type else '-'),
            ('Mode de passation', market.procurement_mode or '-'),
            ('Budget', f"{market.budget:,.2f} MAD" if market.budget else '-'),
            ('Crédits', f"{market.credits:,.2f} MAD" if market.credits else '-'),
            ('Service responsable', market.responsible_service or '-'),
            ('Entreprise attributaire', market.awardee or '-'),
            ('Montant définitif', f"{market.final_amount:,.2f} MAD" if market.final_amount else '-'),
            ('Statut', market.status.value if market.status else '-'),
            ('Date de début', market.start_date.strftime('%d/%m/%Y') if market.start_date else '-'),
            ('Date de fin prévue', market.expected_end_date.strftime('%d/%m/%Y') if market.expected_end_date else '-'),
            ('Date de fin réelle', market.actual_end_date.strftime('%d/%m/%Y') if market.actual_end_date else '-'),
        ]
        
        # Remplir le tableau
        for label, value in info_data:
            row_cells = info_table.add_row().cells
            row_cells[0].text = label
            row_cells[1].text = str(value)
            
            # Style pour les labels
            row_cells[0].paragraphs[0].runs[0].bold = True
        
        # Étapes du marché
        doc.add_page_break()
        doc.add_heading('Étapes du Marché', level=1)
        
        stages = self.db.query(Stage).filter(Stage.market_id == market_id).order_by(Stage.order).all()
        
        if stages:
            stages_table = doc.add_table(rows=1, cols=6)
            stages_table.style = 'Table Grid'
            
            # En-têtes
            headers = stages_table.rows[0].cells
            headers[0].text = 'Ordre'
            headers[1].text = 'Nom'
            headers[2].text = 'Statut'
            headers[3].text = 'Progression'
            headers[4].text = 'Date prévue'
            headers[5].text = 'Date réelle'
            
            # Style des en-têtes
            for cell in headers:
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Données des étapes
            for stage in stages:
                row_cells = stages_table.add_row().cells
                row_cells[0].text = str(stage.order)
                row_cells[1].text = stage.name
                row_cells[2].text = stage.status.value if stage.status else '-'
                row_cells[3].text = f"{stage.progress_percentage}%"
                row_cells[4].text = stage.planned_date.strftime('%d/%m/%Y') if stage.planned_date else '-'
                row_cells[5].text = stage.actual_date.strftime('%d/%m/%Y') if stage.actual_date else '-'
                
                # Colorer les étapes en retard
                if stage.is_late:
                    for cell in row_cells:
                        cell._element.get_or_add_tcPr().append(
                            self._get_shading_element('FFFF00')
                        )
        else:
            doc.add_paragraph('Aucune étape définie pour ce marché.')
        
        # Statistiques
        doc.add_page_break()
        doc.add_heading('Statistiques', level=1)
        
        completed_stages = len([s for s in stages if s.status.value == 'termine'])
        total_stages = len(stages)
        progress = (completed_stages / total_stages * 100) if total_stages > 0 else 0
        
        stats_paragraph = doc.add_paragraph()
        stats_paragraph.add_run(f'Taux de complétion: {progress:.1f}%\n').bold = True
        stats_paragraph.add_run(f'Étapes terminées: {completed_stages}/{total_stages}\n')
        stats_paragraph.add_run(f'Étapes en retard: {len([s for s in stages if s.is_late])}\n')
        
        # Pied de page
        doc.add_page_break()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run(f'Rapport généré le {datetime.now().strftime("%d/%m/%Y %H:%M")}\n')
        footer.add_run('Système de Gestion des Marchés Publics - Communes Territoriales Marocaines')
        
        # Sauvegarder en mémoire
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        
        return output
    
    def create_markets_summary(self, filters: dict = None) -> BytesIO:
        """
        Crée un résumé Word de plusieurs marchés
        
        Args:
            filters: Filtres pour la sélection des marchés
            
        Returns:
            Fichier Word en mémoire
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx n'est pas installé")
        
        query = self.db.query(Market).filter(Market.is_deleted == False)
        
        if filters:
            if 'status' in filters:
                query = query.filter(Market.status == filters['status'])
            if 'year' in filters:
                from sqlalchemy import extract
                query = query.filter(extract('year', Market.created_at) == filters['year'])
        
        markets = query.all()
        
        doc = Document()
        
        # Titre
        title = doc.add_heading('Résumé des Marchés Publics', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date de génération
        doc.add_paragraph(f'Généré le {datetime.now().strftime("%d/%m/%Y %H:%M")}', style='Intense Quote')
        
        # Tableau récapitulatif
        doc.add_heading('Liste des Marchés', level=1)
        
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # En-têtes
        headers = table.rows[0].cells
        headers[0].text = 'Numéro'
        headers[1].text = 'Objet'
        headers[2].text = 'Budget'
        headers[3].text = 'Statut'
        headers[4].text = 'Entreprise'
        
        for cell in headers:
            cell.paragraphs[0].runs[0].bold = True
        
        # Données
        for market in markets:
            row_cells = table.add_row().cells
            row_cells[0].text = market.market_number or '-'
            row_cells[1].text = market.object or '-'
            row_cells[2].text = f"{market.budget:,.2f} MAD" if market.budget else '-'
            row_cells[3].text = market.status.value if market.status else '-'
            row_cells[4].text = market.awardee or '-'
        
        # Statistiques globales
        doc.add_page_break()
        doc.add_heading('Statistiques Globales', level=1)
        
        total_budget = sum(m.budget for m in markets if m.budget)
        total_markets = len(markets)
        
        stats = [
            f'Total des marchés: {total_markets}',
            f'Budget total: {total_budget:,.2f} MAD',
            f'Budget moyen: {(total_budget / total_markets):,.2f} MAD' if total_markets > 0 else 'Budget moyen: -',
        ]
        
        for stat in stats:
            p = doc.add_paragraph()
            p.add_run(stat).bold = True
        
        # Sauvegarder
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        
        return output
    
    def create_analysis_report(self, analysis_data: dict) -> BytesIO:
        """
        Crée un rapport Word pour l'analyse des offres PMMP
        
        Args:
            analysis_data: Données d'analyse du marché
            
        Returns:
            Fichier Word en mémoire
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx n'est pas installé")
        
        doc = Document()
        
        # Titre
        title = doc.add_heading('Rapport d\'Analyse des Offres', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Informations du marché
        doc.add_heading('Informations du Marché', level=1)
        
        market_info = analysis_data.get('market_info', {})
        info_table = doc.add_table(rows=1, cols=2)
        info_table.style = 'Table Grid'
        
        for key, value in market_info.items():
            row_cells = info_table.add_row().cells
            row_cells[0].text = str(key)
            row_cells[1].text = str(value)
            row_cells[0].paragraphs[0].runs[0].bold = True
        
        # Statistiques des offres
        doc.add_page_break()
        doc.add_heading('Statistiques des Offres', level=1)
        
        stats = analysis_data.get('statistics', {})
        stats_table = doc.add_table(rows=1, cols=2)
        stats_table.style = 'Table Grid'
        
        stat_labels = {
            'median': 'Médiane',
            'mean': 'Moyenne',
            'trimmed_mean': 'Moyenne tronquée',
            'min': 'Minimum',
            'max': 'Maximum',
            'std_dev': 'Écart-type'
        }
        
        for key, label in stat_labels.items():
            if key in stats:
                row_cells = stats_table.add_row().cells
                row_cells[0].text = label
                row_cells[1].text = f"{stats[key]:,.2f} MAD" if stats[key] else '-'
                row_cells[0].paragraphs[0].runs[0].bold = True
        
        # Classement des entreprises
        doc.add_page_break()
        doc.add_heading('Classement des Entreprises', level=1)
        
        ranking = analysis_data.get('ranking', [])
        if ranking:
            ranking_table = doc.add_table(rows=1, cols=4)
            ranking_table.style = 'Table Grid'
            
            headers = ranking_table.rows[0].cells
            headers[0].text = 'Rang'
            headers[1].text = 'Entreprise'
            headers[2].text = 'Montant'
            headers[3].text = 'Écart médian'
            
            for cell in headers:
                cell.paragraphs[0].runs[0].bold = True
            
            for company in ranking:
                row_cells = ranking_table.add_row().cells
                row_cells[0].text = str(company.get('rank', '-'))
                row_cells[1].text = company.get('company', '-')
                row_cells[2].text = f"{company.get('amount', 0):,.2f} MAD"
                row_cells[3].text = f"{company.get('percentage_of_median', 0):.1f}%"
        
        # Offres anormales
        abnormal = analysis_data.get('abnormal_offers', {})
        if abnormal and (abnormal.get('low') or abnormal.get('high')):
            doc.add_page_break()
            doc.add_heading('Offres Anormales Détectées', level=1)
            
            if abnormal.get('low'):
                doc.add_heading('Offres anormalement basses', level=2)
                for offer in abnormal['low']:
                    p = doc.add_paragraph()
                    p.add_run(f"{offer['company']}: {offer['amount']:,.2f} MAD").bold = True
                    p.add_run(f" (Écart: {offer['deviation']:.1f}%)")
            
            if abnormal.get('high'):
                doc.add_heading('Offres anormalement élevées', level=2)
                for offer in abnormal['high']:
                    p = doc.add_paragraph()
                    p.add_run(f"{offer['company']}: {offer['amount']:,.2f} MAD").bold = True
                    p.add_run(f" (Écart: {offer['deviation']:.1f}%)")
        
        # Sauvegarder
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        
        return output
    
    def _get_shading_element(self, color: str):
        """
        Crée un élément de shading pour les cellules de tableau
        
        Args:
            color: Couleur hexadécimale
            
        Returns:
            Élément de shading
        """
        from docx.oxml import parse_xml
        shading_xml = f'<w:shd {{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}fill="{color}"/>'
        return parse_xml(shading_xml)


def get_word_export_service(db: Session) -> WordExportService:
    """
    Factory pour créer une instance du service d'export Word
    
    Args:
        db: Session de base de données
        
    Returns:
        Instance de WordExportService
    """
    return WordExportService(db)
